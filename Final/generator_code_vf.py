import requests
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
import getpass

# Función para obtener datos de la API del Banco Mundial
def obtener_datos(api_url):
    print(f"Obteniendo datos de: {api_url}")
    response = requests.get(api_url)
    if response.status_code == 200:
        print("Datos obtenidos con éxito.")
        data = response.json()
        if data and len(data) > 1 and 'value' in data[1][0]:
            # Extraer los valores numéricos, ignorando los valores None
            return [entry['value'] for entry in data[1] if entry['value'] is not None]
        else:
            print("Datos no válidos o vacíos.")
            return []
    


    else:
        print("Error al obtener los datos.")
        return None

# URLs de la API para cada indicador y país
urls = {
    "Gini": {
        "USA": "https://api.worldbank.org/v2/country/USA/indicator/SI.POV.GINI?format=json",
        "India": "https://api.worldbank.org/v2/country/IND/indicator/SI.POV.GINI?format=json",
        "Spain": "https://api.worldbank.org/v2/country/ESP/indicator/SI.POV.GINI?format=json"
    },
    "Desempleo": {
        "USA": "https://api.worldbank.org/v2/country/USA/indicator/SL.UEM.TOTL.ZS?format=json",
        "India": "https://api.worldbank.org/v2/country/IND/indicator/SL.UEM.TOTL.ZS?format=json",
        "Spain": "https://api.worldbank.org/v2/country/ESP/indicator/SL.UEM.TOTL.ZS?format=json"
    },
    "Inflación": {
        "USA": "https://api.worldbank.org/v2/country/USA/indicator/FP.CPI.TOTL.ZG?format=json",
        "India": "https://api.worldbank.org/v2/country/IND/indicator/FP.CPI.TOTL.ZG?format=json",
        "Spain": "https://api.worldbank.org/v2/country/ESP/indicator/FP.CPI.TOTL.ZG?format=json"
    }
}


# Obtener y procesar los datos
datos = {}
for indicador, paises in urls.items():
    datos[indicador] = {}
    for pais, url in paises.items():
        datos[indicador][pais] = obtener_datos(url)

# Crear gráficos
def crear_grafico(datos, indicador):
    print(f"Creando gráfico para {indicador}")
    plt.figure()
    for pais, valores in datos.items():
        print(f"Datos para {pais}: {valores}")  # Seguimiento de datos

        # Generar una lista de los últimos 30 años
        years = list(range(datetime.now().year - 30, datetime.now().year))
        print(f"Años para {pais}: {years}")  # Seguimiento de los años generados

        # Verificar si la longitud de los valores coincide con los años
        if len(valores) >= 30:
            # Tomar solo los últimos 30 valores
            valores_recientes = valores[-30:]
            plt.plot(years, valores_recientes, label=pais)
            plt.xlabel('Año')  # Etiqueta para el eje X
            plt.ylabel(indicador)  # Etiqueta para el eje Y con el nombre del valor representado
        else:
            print(f"Datos insuficientes para {pais}: {valores}")
    # Definiciones de los indicadores
    definiciones = {
        "Gini": "El índice de Gini mide la desigualdad en los ingresos dentro de un país.",
        "Desempleo": "La tasa de desempleo representa el porcentaje de la fuerza laboral que está sin empleo.",
        "Inflación": "La inflación mide el aumento porcentual de los precios al consumidor."
    }

    # Añadir la definición del indicador al gráfico
    # Obtener la definición del indicador actual
    definicion = definiciones.get(indicador, "Definición no disponible.")
    plt.text(0.5, 1.05, definicion, fontsize=10, transform=plt.gca().transAxes, ha='center', va='center', bbox=dict(facecolor='white', alpha=0.5))
    plt.legend()
    plt.savefig(f"{indicador}.png")
    plt.close()

    

for indicador, datos_paises in datos.items():
    crear_grafico(datos_paises, indicador)

# Crear documento Word
print("Creando documento Word...")
doc = Document()
doc.add_heading('Informe Económico 473214', 0)
doc.add_paragraph(
    "Este documento presenta un análisis económico detallado de tres países: Estados Unidos, España e India. El propósito principal de este estudio es proporcionar una visión comparativa de los principales indicadores económicos de cada uno de estos países, con el fin de identificar patrones, similitudes y diferencias en su desempeño económico en un contexto global. La información utilizada en este análisis proviene de la API del Banco Mundial, que es una fuente confiable y actualizada para la obtención de datos económicos globales.\n \n En este análisis se han seleccionado una serie de indicadores clave para evaluar las economías de Estados Unidos, España e India. Estos indicadores incluyen, el índice de Gini, la tasa de inflación y el desempleo. La selección de estos indicadores tiene como objetivo proporcionar una evaluación integral y comparativa de la situación económica-social de los tres países en cuestión."
    "A lo largo del documento, se exploran las variaciones de los últimos 20 años en estos indicadores entre Estados Unidos, España e India, con el fin de identificar tendencias y compararlas en el contexto de sus respectivos marcos económicos, políticas gubernamentales y factores socioeconómicos. Este análisis también tiene la intención de ofrecer una perspectiva clara sobre el estado actual de las economías de estos países y cómo sus políticas pueden estar influyendo en su desarrollo económico.\n \n En resumen, este informe busca ofrecer una visión detallada y comparativa de la situación económica de Estados Unidos, España e India, basada en los datos obtenidos del Banco Mundial, y proporcionar una base sólida para una discusión informada sobre las fortalezas y desafíos económicos de cada nación.\n"
)


# Definiciones de los indicadores
definiciones = {
    "Gini": "El índice de Gini mide la desigualdad en los ingresos dentro de un país. Un valor de 0 representa igualdad perfecta, mientras que un valor de 1 indica desigualdad máxima.",
    "Desempleo": "La tasa de desempleo representa el porcentaje de la fuerza laboral que está sin empleo. Un valor más alto indica un mayor nivel de desempleo.",
    "Inflación": "La inflación mide el aumento porcentual de los precios al consumidor. Un valor más alto indica un aumento más rápido en los precios."
}

# Añadir gráficos al documento con definiciones
for indicador in urls.keys():
    doc.add_heading(indicador, level=1)
    definicion = definiciones.get(indicador, "Definición no disponible.")
    doc.add_paragraph(definicion)
    doc.add_picture(f"{indicador}.png")


# Añadir resumen
doc.add_heading('Resumen Comparativo', level=1)

# Calcular y añadir un resumen de las medias de los indicadores por país
for indicador, datos_paises in datos.items():
    resumen_texto = f"Resumen de {indicador}:\n"
    for pais, valores in datos_paises.items():
        if valores:
            promedio = sum(valores) / len(valores)
            resumen_texto += f"- {pais}: Promedio de los últimos 30 años: {promedio:.2f}\n"
        else:
            resumen_texto += f"- {pais}: Datos insuficientes para calcular el promedio.\n"
    doc.add_paragraph(resumen_texto)

# Añadir valoración de la situación actual de cada país
valoraciones = {
    "USA": "Estados Unidos muestra una economía robusta con desafíos en la desigualdad de ingresos.",
    "India": "India está experimentando un crecimiento económico significativo, aunque enfrenta retos en el desempleo.",
    "Spain": "España ha mostrado mejoras en la inflación, pero el desempleo sigue siendo un desafío importante."
}

for pais, valoracion in valoraciones.items():
    doc.add_paragraph(f"Valoración de {pais}: {valoracion}")



# Añadir imagen con nombre de usuario y fecha/hora
doc.add_heading('Control Usuario', level=1)
usuario = getpass.getuser()
fecha_hora = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
doc.add_paragraph(f"Usuario: {usuario}")
doc.add_paragraph(f"Fecha y hora de ejecución: {fecha_hora}")

# Guardar documento
doc.save('Informe_Economico_473214.docx')
print("Documento guardado como 'Informe_Economico_473214'.")